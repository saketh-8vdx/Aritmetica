from deep_translator import GoogleTranslator
import openai
import boto3
import numpy as np
from langchain.vectorstores import FAISS
from langchain.embeddings import BedrockEmbeddings
from langchain.schema import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
from sklearn.metrics.pairwise import cosine_similarity
from langchain.schema import Document as LangchainDocument
import tiktoken
import json
import streamlit as st
import tempfile
import os
from docx2pdf import convert
from docx import Document
import zipfile
from llama_parse import LlamaParse

encoder = tiktoken.get_encoding("cl100k_base")
MAX_TOKEN_LIMIT = 8000

aws_credentials = st.secrets["aws_credentials"]
bedrock = boto3.client(
    service_name='bedrock-runtime',
    aws_access_key_id=aws_credentials["AWS_ACCESS_KEY_ID"],
    aws_secret_access_key=aws_credentials["AWS_SECRET_ACCESS_KEY"],
    region_name='us-east-1'
)
parser = LlamaParse(
    api_key=st.secrets["LLAMA_KEY"]["llama_key"],  # Add your API key here if needed
    result_type="markdown",
    verbose=True,
    language="es"
)

openai.api_key = st.secrets["openai"]["api_key"]


# class CustomBedrockEmbeddings(BedrockEmbeddings):
#
#     def embed_documents(self, texts):
#         # Use concurrent processing to embed multiple documents at the same time
#         with ThreadPoolExecutor() as executor:
#             futures = [executor.submit(self.embed_query, text) for text in texts]
#             results = [future.result() for future in as_completed(futures)]
#         return results
#
#     def embed_query(self, text):
#         body = json.dumps({
#             "inputText": text,
#             "dimensions": 1024,
#             "normalize": True
#         })
#
#         response = bedrock.invoke_model(
#             body=body,
#             modelId='amazon.titan-embed-text-v2:0',
#             contentType='application/json',
#             accept='application/json'
#         )
#
#         response_body = json.loads(response['body'].read())
#         return response_body['embedding']
#
#
# embeddings = CustomBedrockEmbeddings()


class CustomBedrockEmbeddings(BedrockEmbeddings):
    def __init__(self, client=None, **kwargs):
        # Pass credentials through super().__init__
        super().__init__(
            client=client,
            region_name="us-east-1",
            model_id="amazon.titan-embed-text-v2:0",
            **kwargs
        )

    def embed_documents(self, texts):
        # Use concurrent processing to embed multiple documents at the same time
        with ThreadPoolExecutor() as executor:
            futures = [executor.submit(self.embed_query, text) for text in texts]
            results = [future.result() for future in as_completed(futures)]
        return results

    def embed_query(self, text):
        # Ensure text is a string and not empty
        if not isinstance(text, str) or not text.strip():
            return [0] * 1024  # Return zero vector for empty text

        body = json.dumps({
            "inputText": text,
            "dimensions": 1024,
            "normalize": True
        })

        try:
            response = self.client.invoke_model(
                body=body,
                modelId='amazon.titan-embed-text-v2:0',
                contentType='application/json',
                accept='application/json'
            )

            response_body = json.loads(response['body'].read())
            return response_body['embedding']
        except Exception as e:
            st.error(f"Error during embedding: {str(e)}")
            # Return zero vector in case of error
            return [0] * 1024


# Initialize embeddings with the bedrock client
embeddings = CustomBedrockEmbeddings(client=bedrock)


def split_text_with_token_limit(documents, max_tokens=8000, max_chars=49000):
    """
    Splits documents into chunks that respect both token and character limits.
    Args:
        documents: List of document objects
        max_tokens: Maximum number of tokens per chunk
        max_chars: Maximum characters per chunk (AWS Bedrock limit)
    """
    chunked_documents = []

    for document in documents:
        text = document.page_content
        # First split by paragraphs
        paragraphs = text.split('\n')

        current_chunk = ""
        char_chunks = []

        # First pass: split by character limit
        for para in paragraphs:
            if len(current_chunk) + len(para) > max_chars:
                if current_chunk:
                    char_chunks.append(current_chunk)
                current_chunk = para
            else:
                current_chunk += '\n' + para if current_chunk else para

        if current_chunk:
            char_chunks.append(current_chunk)

        # Second pass: split by tokens
        for char_chunk in char_chunks:
            tokens = encoder.encode(char_chunk)
            current_chunk = []
            current_token_count = 0

            for token in tokens:
                if current_token_count + 1 > max_tokens:
                    chunk_text = encoder.decode(current_chunk)
                    if len(chunk_text) > 0:  # Only add non-empty chunks
                        chunked_documents.append(
                            LangchainDocument(
                                page_content=chunk_text,
                                metadata=document.metadata
                            )
                        )
                    current_chunk = [token]
                    current_token_count = 1
                else:
                    current_chunk.append(token)
                    current_token_count += 1

            if current_chunk:
                chunk_text = encoder.decode(current_chunk)
                if len(chunk_text) > 0:  # Only add non-empty chunks
                    chunked_documents.append(
                        LangchainDocument(
                            page_content=chunk_text,
                            metadata=document.metadata
                        )
                    )

    print(f"Number of chunks: {len(chunked_documents)}")
    return chunked_documents


def split_text_with_token_limit_new(documents, max_tokens=8000, max_chars=49000):
    """
    Splits documents into chunks that ensure exactly 8000 tokens while staying under the character limit.
    Args:
        documents: List of document objects
        max_tokens: Maximum number of tokens per chunk (default: 8000)
        max_chars: Maximum characters per chunk (AWS Bedrock limit: 49,000)
    """
    chunked_documents = []

    for document in documents:
        text = document.page_content

        # Tokenize the entire document at once
        tokens = encoder.encode(text)
        total_tokens = len(tokens)

        start_idx = 0

        while start_idx < total_tokens:
            # Extract exactly `max_tokens` tokens (or fewer if at the end)
            token_chunk = tokens[start_idx: start_idx + max_tokens]
            chunk_text = encoder.decode(token_chunk)

            # Ensure the chunk does not exceed `max_chars`
            while len(chunk_text) > max_chars:
                token_chunk = token_chunk[:-1]  # Remove one token at a time
                chunk_text = encoder.decode(token_chunk)

            # Store the chunk
            chunked_documents.append(
                LangchainDocument(
                    page_content=chunk_text,
                    metadata=document.metadata
                )
            )

            # Move to the next batch of tokens
            start_idx += len(token_chunk)

    print(f"Number of chunks: {len(chunked_documents)}")
    return chunked_documents


def process_markdown_document(doc):
    """Helper function to process a single document."""
    markdown_text = doc.get_content()
    print("Original Markdown Text:", markdown_text)

    # Translate to English
    translated_text = GoogleTranslator(source='auto', target='en').translate(markdown_text)
    print("Translated Text:", translated_text)

    return translated_text


def process_documents_and_translate(documents):
    translated_texts = []

    # Use ThreadPoolExecutor for parallel processing
    with ThreadPoolExecutor() as executor:
        futures = [executor.submit(process_markdown_document, doc) for doc in documents]

        for future in as_completed(futures):
            try:
                translated_texts.append(future.result())
            except Exception as e:
                print(f"Error processing document: {e}")

    print("Final Translated Texts:", translated_texts)
    return translated_texts


def extract_zip_files(folder_path, extract_to):
    """
    Extracts all ZIP files in the given folder to the specified directory.
    """
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".zip"):
                zip_path = os.path.join(root, file)
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(extract_to)


def translate_text(text):
    """Helper function to process a single document."""
    markdown_text = text
    print("Original Markdown Text:", markdown_text)

    # Translate to English
    translated_text = GoogleTranslator(source='auto', target='en').translate(markdown_text)
    print("Translated Text:", translated_text)

    return translated_text


def read_docx(file_path):
    """
    Reads content from a DOCX file and returns translated text.
    """
    doc = Document(file_path)
    texts = []
    for paragraph in doc.paragraphs:
        print(paragraph.text)
        texts.append(paragraph.text)
    # texts = [para.text for para in doc.paragraphs if para.text.strip()]

    # Use parallel processing for translation
    with ThreadPoolExecutor() as executor:
        translated_texts = list(executor.map(translate_text, texts))

    return translated_texts


def find_files(folder_path):
    matching_files = []
    all_chunks = []
    i = 0

    extract_to = os.path.join(folder_path, "extracted")
    os.makedirs(extract_to, exist_ok=True)

    try:
        extract_zip_files(folder_path, extract_to)
    except Exception as e:
        print(f"Error extracting zip files: {e}")

    for search_path in [folder_path, extract_to]:
        for root, _, files in os.walk(search_path):
            for file in files:
                try:
                    chunked_documents = []
                    file_path = os.path.join(root, file)

                    if file.lower().endswith(".pdf"):
                        try:
                            print("pdf")
                            print(file)
                            documents = parser.load_data(file_path)
                            texts = process_documents_and_translate(documents)
                            text_content = "\n".join(texts)
                            document = LangchainDocument(page_content=text_content)
                            chunked_documents = split_text_with_token_limit_new([document], max_tokens=8000)
                        except Exception as e:
                            print(f"Error processing PDF {file}: {e}")

                    elif file.lower().endswith(".docx"):
                        try:
                            print("docx")
                            print(file)
                            convert(file_path, f"output_{i}.pdf")
                            file_path = f"output_{i}.pdf"
                            documents = parser.load_data(file_path)
                            texts = process_documents_and_translate(documents)
                            text_content = "\n".join(texts)
                            document = LangchainDocument(page_content=text_content)
                            chunked_documents = split_text_with_token_limit_new([document], max_tokens=8000)
                        except Exception as e:
                            print(f"Error processing DOCX {file}: {e}")

                    i += 1
                    print(f"value of i = {i}")
                    all_chunks.extend(chunked_documents)
                except Exception as e:
                    print(f"Error processing file {file}: {e}")

    print(len(all_chunks))
    return all_chunks


def store_embeddings_faiss(chunked_documents, faiss_index_path):
    vectorstore = FAISS.from_documents(chunked_documents, embeddings)
    vectorstore.save_local(faiss_index_path)


def rerank_documents(query, documents):
    query_embedding = np.array(embeddings.embed_query(query)).reshape(1, -1)  # Convert to 2D array

    similarities = []
    for doc in documents:
        doc_embedding = np.array(embeddings.embed_query(doc.page_content)).reshape(1, -1)  # Convert to 2D array

        similarity = cosine_similarity(query_embedding, doc_embedding)[0][0]  # Extract the similarity score
        similarities.append((doc, similarity))

    ranked_docs = [doc for doc, _ in sorted(similarities, key=lambda item: item[1], reverse=True)]
    return ranked_docs


def generate_response_temp(query, retriever):
    try:
        docs = retriever.get_relevant_documents(query)
        rerank_docs = rerank_documents(query, docs)
        rerank_docs = rerank_docs[:min(30, len(rerank_docs))]
        context = "\n".join([doc.page_content for doc in rerank_docs])

        # context = "\n\n".join([doc.page_content for doc in docs])

        messages = [
            {"role": "system",
             "content": """You are a specialized assistant for analyzing Legal reports. Your task is to generate response
             based on the input question from the given context with high accuracy.Understand the Questions and answer it accurately from the given context.
             Give the proper explanation for each response"""},
            {"role": "user", "content": f"Context: {context}\n\nQuestion: {query}\n\nAnswer:"}
        ]

        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=messages,
            temperature=0.1,
            top_p=1,
            frequency_penalty=0.0,
            presence_penalty=0.0
        )
        # message = response.choices[0].message
        print(response)

        final_response = response.choices[0].message.content.strip()
        return final_response


    except Exception as e:
        print(f"generate response error: {e}")


# Streamlit UI
st.set_page_config(page_title="Legal Document Analyzer", layout="wide")

st.title("Legal Document Analyzer")
st.write("Upload documents and ask questions to get insights from your legal documents.")

# Sidebar for uploading files
with st.sidebar:
    st.header("Upload Documents")
    uploaded_files = st.file_uploader("Choose files", accept_multiple_files=True, type=["pdf", "docx", "zip"])

    process_button = st.button("Process Documents")

    st.header("Settings")
    faiss_index_name = st.text_input("FAISS Index Name", "legal_documents_index")

# Main area for asking questions
col1, col2 = st.columns([2, 1])

with col1:
    st.header("Ask Questions")
    query = st.text_area("Enter your question:", height=100)
    submit_button = st.button("Submit Question")

# Process uploaded files
if process_button and uploaded_files:
    with st.spinner("Processing documents..."):
        # Create a temporary directory to store uploaded files
        temp_dir = tempfile.mkdtemp()

        # Save uploaded files to the temporary directory
        for uploaded_file in uploaded_files:
            file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

        # Process the files
        st.info(f"Processing files from {temp_dir}...")

        chunks = find_files(temp_dir)

        if chunks:
            # Create FAISS index directory if it doesn't exist
            os.makedirs("faiss_indexes", exist_ok=True)
            faiss_index_path = os.path.join("faiss_indexes", faiss_index_name)

            # Store embeddings
            store_embeddings_faiss(chunks, faiss_index_path)
            vectorstore = FAISS.load_local(faiss_index_path, embeddings, allow_dangerous_deserialization=True)

            st.session_state.vectorstore = vectorstore
            st.session_state.faiss_index_path = faiss_index_path

            st.success(f"Documents processed successfully !")
            st.info(f"Created {len(chunks)} chunks from the documents.")
        else:
            st.error("No documents were processed. Please check the file formats.")

# Generate response
if submit_button and query:
    if not hasattr(st.session_state, 'vectorstore'):
        if os.path.exists(os.path.join("faiss_indexes", faiss_index_name)):
            st.session_state.faiss_index_path = os.path.join("faiss_indexes", faiss_index_name)
            st.session_state.vectorstore = FAISS.load_local(
                st.session_state.faiss_index_path,
                embeddings,
                allow_dangerous_deserialization=True
            )
        else:
            st.error("Please process documents first or ensure that the FAISS index exists.")
            st.stop()

    with st.spinner("Generating response..."):
        retriever = st.session_state.vectorstore.as_retriever(search_kwargs={"k": 35})
        response = generate_response_temp(query, retriever)

        with col2:
            st.header("Response")
            st.write(response)

# Display information about the current state
with st.sidebar:
    if hasattr(st.session_state, 'faiss_index_path'):
        st.success(f"Using FAISS index: {st.session_state.faiss_index_path}")
    else:
        st.info("No FAISS index loaded yet. Process documents first.")
